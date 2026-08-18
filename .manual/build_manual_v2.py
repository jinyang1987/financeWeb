# -*- coding: utf-8 -*-
"""
生成《会计档案管理系统操作手册》docx
- 正文：宋体四号字（14pt）、行间距 1.5 倍、首行缩进 2 字符
- 标题：黑体加粗（一号/三号分级）
- 表格：宋体小五，带边框
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = Document()

# ============ 全局页面设置 ============
for section in DOC.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)

# ============ 工具函数 ============

def set_font(run, name_cn="宋体", name_en="Times New Roman", size=14,
             bold=False, color=None):
    run.font.name = name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text, *, size=14, bold=False, indent=True, align=None,
             space_before=0, space_after=6, line=1.5, font_cn="宋体",
             color=None, keep_with_next=False):
    """普通正文段落：宋体四号(14)、行距1.5倍、首行缩进2字符"""
    p = DOC.add_paragraph()
    pf = p.paragraph_format
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = line
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        # 首行缩进 2 字符（用 firstLineChars 实现按字符缩进）
        pPr = p._p.get_or_add_pPr()
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        ind.set(qn('w:firstLineChars'), '200')
        ind.set(qn('w:firstLine'), '480')  # 2字符 * 240 twips
    if align == 'center':
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if keep_with_next:
        pf.keep_with_next = True
    run = p.add_run(text)
    set_font(run, name_cn=font_cn, size=size, bold=bold, color=color)
    return p

def add_h1(text):
    """一级标题：黑体 三号(16pt) 加粗，蓝色"""
    return add_para(text, size=16, bold=True, indent=False,
                    space_before=18, space_after=10, font_cn="黑体",
                    color=(0x0F, 0x3B, 0x63), keep_with_next=True)

def add_h2(text):
    """二级标题：黑体 四号(14pt) 加粗"""
    return add_para(text, size=14, bold=True, indent=False,
                    space_before=12, space_after=8, font_cn="黑体",
                    color=(0x15, 0x5E, 0x9C), keep_with_next=True)

def add_h3(text):
    """三级标题：宋体 四号加粗"""
    return add_para(text, size=14, bold=True, indent=False,
                    space_before=8, space_after=6, font_cn="黑体",
                    color=(0x33, 0x33, 0x33), keep_with_next=True)

def add_bullet(text, level=0):
    """列表项：宋体四号，1.5倍行距，项目符号，左缩进"""
    p = DOC.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.8 + 0.6 * level)
    pf.first_line_indent = Cm(-0.4)
    run = p.add_run("•  " + text)
    set_font(run, size=14)
    return p

def add_num(n, text):
    """带序号列表项"""
    p = DOC.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.8)
    pf.first_line_indent = Cm(-0.4)
    run = p.add_run(f"{n}. {text}")
    set_font(run, size=14)
    return p

def set_cell(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.2
    pf.space_after = Pt(2)
    pf.space_before = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)

def add_table(headers, rows, col_widths=None):
    """表格：表头加粗灰底，内容宋体小五"""
    t = DOC.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8F0FE')
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    # 数据行
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], str(val))
    # 列宽
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    # 表后空行
    add_para("", size=6, indent=False, space_after=2)
    return t

# ============ 封面 ============
add_para("", indent=False, space_after=0)
add_para("", indent=False, space_after=0)
add_para("会计档案管理系统", size=26, bold=True, indent=False,
         align='center', font_cn="黑体", color=(0x0F, 0x3B, 0x63), space_before=20)
add_para("操 作 手 册", size=22, bold=True, indent=False,
         align='center', font_cn="黑体", color=(0x0F, 0x3B, 0x63), space_after=10)
add_para("（电子会计档案全生命周期管理平台 · 基于 DA/T 13/42/94 与《会计档案管理办法》79 号令设计）",
         size=12, indent=False, align='center', font_cn="宋体",
         color=(0x64, 0x74, 0x8B), space_after=6)
add_para("版本 V2.0 · 2026-08-18", size=12, indent=False,
         align='center', font_cn="宋体", color=(0x64, 0x74, 0x8B), space_after=30)

add_para("编制说明", size=14, bold=True, indent=False, font_cn="黑体",
         space_before=6, keep_with_next=True)
add_para("本手册依据系统最新代码（financeWeb 前端 + ams-server 业务服务 + Alfresco 内容库 + PostgreSQL 数据库，"
         "混合架构全链路真贯通）重编，全面覆盖系统概述、角色与登录、检索门户、档案收集、档案整理、档案保管、"
         "档案利用、档案处置、档案统计、档案配置、系统管理等全部业务域，并对核心状态机、档号规则、检测项、"
         "常见问题与服务运维作了系统化说明。每个功能按「用途 / 入口 / 操作步骤 / 规则与注意事项」组织，"
         "便于读者按需定位、照单操作。")

add_para("系统主线：档案收集（抓取 / 推送 / 上传）→ 核对审核 → 组卷 → 四性检测 → 确认赋号 → 移交归盒 → "
         "上架保管 → 检索利用（借阅审批）→ 到期鉴定销毁，全流程状态机驱动、操作留痕（哈希链防篡改）。",
         space_after=12)

DOC.add_page_break()

# ============ 目录占位（正文） ============

# ============ 第 1 章 系统概述 ============
add_h1("第 1 章 系统概述")

add_h2("1.1 系统定位")
add_para("会计档案管理系统是面向财务与档案部门的电子会计档案全生命周期管理平台，覆盖会计凭证、会计账簿、"
         "财务会计报告、其他会计资料四大门类，同时支持纯电子档案与纸质档案数字化双模式。系统依据国家档案局 "
         "DA/T 13-2022《档号编制规则》、DA/T 42-2009《企业文件材料归档范围和档案保管期限规定》、"
         "DA/T 94-2022《电子文件归档与电子档案管理规范》以及财政部、国家档案局《会计档案管理办法》（79 号令）"
         "设计实现，满足档案业务全过程的规范化管理要求。")

add_h2("1.2 系统架构")
add_para("系统采用「混合架构」运行，遵循「Alfresco 管内容、ams 管过程、PG 统一落库」的分工原则：")
add_bullet("Alfresco Content Services（开源内容管理平台）：负责文件二进制、全宗、目录树、件（档案）、卷、盒、"
           "原始凭证子件等节点与自定义属性（finance-model 自定义模型）的存储，同时承载组织、人员、角色等 "
           "Alfresco Groups / People 基础实体与 Ticket 认证。")
add_bullet("ams-server 业务服务（Spring Boot）：承载业务过程数据与规则引擎，包括档号流水（原子取号）、借阅四表、"
           "审批、履约、操作日志（哈希链防篡改）、四性检测报告、封装包、移交批次、鉴定销毁、库房树、配置中心、"
           "用友同步、开放推送、收集台账等，全部真持久化。")
add_bullet("PostgreSQL 数据库：同一实例内分为 schema alfresco（Alfresco 自身表，内容侧）与 schema ams "
           "（21 张业务表，Flyway V1–V5 管理，业务过程侧）。")
add_bullet("前端 financeWeb（React 19 + Vite + TypeScript + Zustand）：提供「检索门户（前台）」与「后台管理」"
           "两套工作区，通过 /api/ams 代理调用 ams-server、/api/proxy/alfresco 调用 Alfresco。")
add_para("配置中心统一由 ams_config 一张 KV 表（jsonb）承载全部前端配置域，前端 configStorage.ts 适配器统一"
         "读写，API 不可达时自动降级 localStorage，保证离线可演示。")

add_h2("1.3 业务主线与数据流")
add_table(
    ["环节", "入口菜单", "产出"],
    [
        ["收集", "档案收集：抓取收集中台 / 集成接口采集；档案整理：核对工作台", "收集池待办件（仅件数据）+ 收集台账"],
        ["整理", "档案整理：核对工作台 / 组卷工作台", "案卷（草稿 → 已确认，赋档号）"],
        ["保管", "档案保管：财务分类视图 / 实体档案库房", "盒（装盒 → 封盒 → 上架在架）"],
        ["利用", "检索门户 / 档案利用：审批中心、借阅管理、借阅台账", "借阅单与履约记录（七态主单 + 履约十态）"],
        ["处置", "档案处置：档案打包 / 档案移交 / 鉴定销毁", "移交批次、销毁留痕"],
    ],
    col_widths=[2.0, 6.5, 5.0]
)

add_h2("1.4 角色体系")
add_table(
    ["角色", "职责", "典型菜单"],
    [
        ["普通员工", "检索与发起借阅", "检索门户"],
        ["部门经理", "借阅一级审批", "检索门户、审批中心"],
        ["档案管理员", "收集 / 整理 / 保管 / 利用 / 配置全流程", "全部业务菜单"],
        ["档案主管", "库房与鉴定销毁管理", "业务菜单（无配置组）"],
        ["财务总监", "大额 / 高危权限审批", "查询、审批中心、台账"],
        ["HR 副总裁", "涉密档案审批", "查询、审批中心"],
        ["系统管理员", "全部功能与系统管理", "全部菜单"],
    ],
    col_widths=[2.5, 6.5, 4.5]
)

# ============ 第 2 章 登录与界面导航 ============
add_h1("第 2 章 登录与界面导航")

add_h2("2.1 登录与身份切换")
add_para("用途：进入系统并确定当前工作身份。", space_after=3)
add_num(1, "打开系统首页，在登录页点选演示账号（各账号预置不同角色，系统内置 7 个演示账号 / 7 个角色组）。")
add_num(2, "登录后默认进入「统计驾驶舱」（有后台权限的角色）或检索门户（普通员工）。")
add_num(3, "如需切换身份，点击页头右上角头像，在「切换身份（演示）」下拉中选择其他账号，无需重新登录。")
add_para("安全说明：系统认证走 Alfresco Ticket 真实校验，ams-server 内存会话有效期 12 小时；"
         "借阅审批、出库、归还、中止等操作均按服务端会话严格校验角色与身份，防止越权与冒名。")

add_h2("2.2 前台与后台")
add_para("系统分「检索门户」（前台）与「后台管理」两个工作区，前后台共用登录态与数据源，可随时互切：")
add_bullet("页头右上角「检索门户」按钮：后台 → 前台。")
add_bullet("门户顶栏「进入后台管理」按钮：前台 → 后台（仅拥有后台菜单权限的角色显示）。")

add_h2("2.3 页头与菜单")
add_bullet("页头左侧：当前位置（菜单组 / 页面名）；右侧：检索门户入口、当前全宗选择器、身份切换、退出。")
add_bullet("左侧菜单按业务域分组：档案查询 / 档案收集 / 档案整理 / 档案保管 / 档案利用 / 档案统计 / "
           "档案处置 / 档案配置 / 系统管理；菜单可见性由角色决定（角色管理可配）。")
add_bullet("切换全宗后，各列表与统计自动按全宗过滤；多全宗单位注意先选对全宗再操作。")

# ============ 第 3 章 检索门户（前台） ============
add_h1("第 3 章 检索门户（前台）")

add_h2("3.1 门户首页")
add_para("用途：面向全体员工的百度式检索入口。检索门户所有登录用户均可使用。", space_after=3)
add_num(1, "在中央大搜索框输入凭证号 / 摘要 / 往来单位 / 单据号 / 档号，点击「检索」进入综合检索结果页。")
add_num(2, "也可用「快捷分类」（全部 / 会计凭证 / 会计账簿 / 财务报表 / 其他资料）一键限定类别。")
add_num(3, "下方六张「检索能力」卡片（综合 / 凭证 / 事项 / 附件 / 关联 / 审计）可直达对应检索模式。")
add_num(4, "首页展示数据概览（如已归档凭证数，由服务端 SQL 统计，不占用全量拉取）。")

add_h2("3.2 六种检索模式")
add_table(
    ["模式", "适用场景", "结果展示"],
    [
        ["综合检索", "全库关键词 + 类别 / 年度 / 科目 / 部门 / 金额高级筛选", "表格：档号 / 凭证号 / 摘要 / 类别 / 期间 / 科目 / 部门 / 金额 / 状态"],
        ["凭证检索", "凭证号 / 科目 / 年度 / 主体 / 制单人 / 金额组合查询", "表格：凭证号 / 摘要 / 科目 / 期间 / 制单人 / 金额 / 状态"],
        ["事项检索", "按经济业务定位：往来单位 / 发票号 / 业务类型", "表格：单据编号 / 类型 / 业务 / 日期 / 单位 / 金额 / 所属凭证"],
        ["附件检索", "原始凭证附件：类型树 + 载体 / 四性 / 日期 / 金额", "左类型树 + 右表格，父件信息联动"],
        ["关联查询", "纸质副本与原生电子同屏比对（卷 / 件 / 盒定位）", "元数据对比表"],
        ["审计追踪", "操作日志哈希链穿透（可导出取证包）", "时间链 + 哈希校验结果"],
    ],
    col_widths=[2.2, 6.3, 5.0]
)
add_para("说明：检索数据源为 ams-server 真后端全量件口径（收集池 ∪ 案卷库 ∪ 盒库），"
         "与后台「档案查询」二级菜单数据完全一致，已归档件均可正常检索定位。")

add_h2("3.3 结果与详情")
add_num(1, "结果表格支持分页（默认 20 条 / 页，可改 10 / 50 / 100）。")
add_num(2, "点击行进入档案详情：元数据、卷级信息（卷号 / 盒号归属）、附件清单。")
add_num(3, "「电子版可用 / 实体在库 / 实体借出」状态标签直接决定能否在线调阅或需申请借阅。")
add_para("附件权限门控：当前用户仅当存在「生效中的电子授权」时方可预览 / 下载原始凭证附件；"
         "未授权仅展示附件清单，并引导「加入借阅车 → 去结算」申请借阅。")

add_h2("3.4 我的借阅（门户）")
add_para("门户顶栏「我的借阅」与后台借阅数据一致，含借阅车、我的申请（审批进度）、在线调阅（限时授权）三部分，"
         "详见第 7 章。")

# ============ 第 4 章 档案收集 ============
add_h1("第 4 章 档案收集")

add_h2("4.1 抓取收集中台（菜单：档案收集 → 抓取收集中台）")
add_para("用途：从已配置的会计核算系统（如用友 BIP）按会计期间主动抓取凭证并归档转换。"
         "用友集成已真实打通 7 个接口（获取 token、账簿查询、期间查询、凭证列表、凭证详情、凭证附件下载、"
         "报表查询），支持定时调度与手动同步。", space_after=3)
add_num(1, "① 选择数据源：只读展示连接状态；数据源连接统一在「系统管理 → 连接配置 → 数据源连接」维护（多数据源："
           "用友 / 金蝶 / 发票 / 银行 / 报销等），抓取人员仅负责抓取与查看日志，无权改动配置。")
add_num(2, "② 选择会计期间：下拉为真实期间（来自用友期间接口），旁边实时显示该期间用友侧凭证数。")
add_num(3, "③ 选择抓取去向（本次抓取的数据流向，默认值在连接配置中维护）：")
add_bullet("直接入库 · 自动组卷（auto-archive）：建件 → 四性检测 → 按类别 + 保管期限自动建卷 → 确认取号归档（可信源适用）。")
add_bullet("送组卷工作台（to-volume）：进入待组卷池，人工组卷。")
add_bullet("送核对工作台 · 待核对（to-check）：登记收集台账，先核对凭证连续性 / 附件，通过后送组卷或转审核。")
add_bullet("送核对工作台 · 待审核（to-review）：跳过核对直接人工审核，通过后送组卷。")
add_num(4, "点击「立即抓取」；下方批次历史表可展开每批明细（类型 / 凭证字号 / 金额 / 状态 / 档号）。")
add_para("规则与注意：以用友凭证 ID 幂等去重，重复抓取同一期间自动跳过；无电子附件的凭证按版式自动生成 PDF "
         "作为电子文件（符合 79 号令）；抓取成功后前端自动刷新件域镜像。调度配置（启用 / cron / 默认去向）可在"
         "「连接配置」中设置，支持定时自动抓取。")

add_h2("4.2 集成接口采集（菜单：档案收集 → 集成接口采集）")
add_para("用途：业务系统按统一四类契约（凭证 / 账簿 / 报告 / 其他）向档案系统推送数据，实现跨系统电子档案自动采集。", space_after=3)
add_bullet("统一推送契约 v2：批次级字段（period / category / destination / runFourChecks / items[]），"
           "条目公共字段 + 四类特有块（凭证含分录、账簿含类型与科目、报表含名称期间、其他含资料类型）。")
add_bullet("推送监控：批次总览、四性检测、去向操作（送审核 / 自动组卷）、批次明细、全链路推送日志（ams_push_log）。")
add_bullet("接口标准：契约字段、样例报文、保管期限对照表，供推送方开发参考（支持 OFD / PDF / A 格式文件 base64 上传）。")
add_bullet("模拟推送：内置 demo-simulator 应用，按类别生成仿真样例，走真实 pushBatch 管道，批次 / 四性 / 日志 / "
           "去向全部真实可演示，后续接入真系统零改动。")
add_para("保管期限缺省自动带出：凭证 / 账簿 30 年、年报永久、中期报告 10 年、其他 10 年（与三合一表口径一致）。"
         "推送方使用签发的 AppKey / AppSecret 做 Bearer 认证，externalId + sourceSystem 幂等去重。")

add_h2("4.3 核对工作台（菜单：档案整理 → 核对工作台）")
add_para("用途：组卷前唯一准入关口（数据质量视角的「核对」与准入审批视角的「审核」已合并为一页），三个 Tab：", space_after=3)
add_bullet("待核对：凭证号连续性核对 + 补传附件（扫描件，真 multipart 上传）+ 收集池待核对（to-check 去向，"
           "通过可「送组卷」或「转审核」）。")
add_bullet("待审核：抓取 / 推送 to-review 去向的数据，人工审核通过 / 驳回（驳回必填意见），留痕于 ams_review_log。")
add_bullet("已处理：审核通过 / 驳回的历史留痕（后端 GET /review/processed 真实数据）。")

# ============ 第 5 章 档案整理（组卷工作台） ============
add_h1("第 5 章 档案整理（组卷工作台）")
add_para("用途：把收集池中的件组成案卷，完成排序、四性检测、赋号、移交。菜单：档案整理 → 组卷工作台。"
         "页面左为待组卷池（排除收集台账未核对件），右为案卷列表。")

add_h2("5.1 组卷（建卷与加件）")
add_num(1, "在左侧勾选凭证（可按年 / 月 / 类别 / 期限筛选，顶部有凭证号连续性提示条与「N 件未核对」提示）。")
add_num(2, "点「组卷（N 件）」一键建卷并加入；也可「加入已有」到既有草稿卷，或点案卷卡片上「加入当前案卷」。")
add_num(3, "「智能组卷」按组卷盒号配置自动推荐分组（按类别 / 期间 / 连续号段），可整组接受或取消。")
add_num(4, "案卷列表右上角「新建案卷」可建空卷（沿用当前筛选的类别 / 期限），改名后从左侧加件。")

add_h2("5.2 卷内件操作（勾选后使用底部悬浮工具栏）")
add_para("在案卷卡片中勾选卷内件前的复选框，页面底部会弹出深色悬浮工具栏：")
add_bullet("全选：选中本卷全部件；上移 / 下移：调整选中件的卷内顺序（可连击逐位移动）。")
add_bullet("转卷：选中件移入其他草稿卷（类别 / 年度 / 期限须一致）。")
add_bullet("拆分：选中件拆出为新案卷（继承本卷属性；全部拆出时本卷自动销毁）。")
add_bullet("移出回池：选中件回到左侧待组卷池。")

add_h2("5.3 卷级操作")
add_bullet("合并：将其他同类别 / 年度 / 期限的草稿卷并入本卷（来源卷删除）。")
add_bullet("拆卷：整卷打散，全部件回待组卷池，案卷删除（有确认弹窗）。")
add_bullet("目录预览 / 打印：卷内目录打印。")
add_bullet("删除空卷：空草稿卷可直接删除。")

add_h2("5.4 四性检测与确认组卷")
add_num(1, "点「运行四性检测」：服务端按检测方案（inspection.plan）对卷内每件执行已启用的归档环节检测项，"
           "并做卷级检测：凭证号断号、卷内查重、件数一致。")
add_num(2, "四个维度（真实性 / 完整性 / 可用性 / 安全性）全部通过，「确认组卷」才可点；未通过时卡片内列出问题明细，处理后重新检测。")
add_num(3, "「确认组卷」：按档号规则配置赋号（组卷时赋号）或仅确认不赋号（会计档案自有凭证号体系），"
           "档号由后端 ams_code_serial 乐观锁原子取号（绝不重号），案卷状态 → 已确认。")

add_h2("5.5 移交与退回")
add_bullet("「移交至档案保管」：已确认案卷自动归入对应类别 / 年度的档案盒（无活动盒自动建盒），状态 → 已移交。")
add_bullet("「撤销确认」：已确认案卷回草稿（档号回收为占位，重新确认时再取号）；已移交案卷可在盒管理中「退回组卷工作台」。")
add_para("规则与注意：拆分 / 合并 / 转卷仅限草稿卷且要求同类别 / 年度 / 保管期限；已确认卷须先撤销确认；"
         "空卷自动销毁。")

# ============ 第 6 章 档案保管 ============
add_h1("第 6 章 档案保管")

add_h2("6.1 财务分类视图（菜单：档案保管 → 财务分类视图）")
add_para("用途：按「年度 → 类别 → 案卷盒 → 卷内件」层级浏览已归档电子会计档案（盒 → 卷 → 件只读投影，"
         "全量明细台账），支持筛选、详情、导出。仅展示已移交（transferred）案卷，形成闭环。")

add_h2("6.2 实体档案库房（菜单：档案保管 → 实体档案库房）")
add_para("用途：实体档案的密集架可视化保管。移交归盒后的盒在「待上架区」排队，上架后进入密集架在架管理。"
         "库房数据由后端 /storage/tree（库房树）与 /storage/occupancy（占用统计）真实驱动，盒上架写真实端点。", space_after=3)
add_para("密集架阵列操作：")
add_num(1, "每架默认闭合（薄板，占用量自下而上染色）；点击某列「打开通道」查看该列各层盒位（每架同时只开一列，还原真实密集架）。")
add_num(2, "彩色盒脊按类别区分（凭证蓝 / 账簿紫 / 报表绿 / 其他琥珀；灰色为他全宗盒）；点击盒脊查看盒详情。")
add_para("上架操作：")
add_bullet("自动上架：服务端按「架 → 列 → 层 → 位」顺序分配第一个空格位，落点列自动打开并高亮。")
add_bullet("点选架位：进入放置模式后点击任意空格位上架；ESC 退出。")
add_bullet("换架位：在架盒详情中点「换架位」，点选新格位后原格位自动释放。")
add_bullet("下架：在架盒下架回到「已封盒」待处理，格位释放。")
add_para("盒状态机：装盒中（active）→ 已封盒（sealed）→ 在架（stored）→ 下架回已封盒；在架盒不可开封 / 删除；"
         "封盒 / 开封 / 上架 / 删除（空盒守卫）均走后端写端点。库房布局在「档案配置 → 库房配置」维护（第 10.8 节）。")

add_h2("6.3 盒与卷的关系")
add_para("盒是卷的物理容器（盒 → 卷 → 件三级）。移交归盒自动进行；盒满后「封盒」不再接收新卷；"
         "在库统计与盒详情在库房页与财务分类视图中均可查。")

# ============ 第 7 章 档案利用（借阅与审批） ============
add_h1("第 7 章 档案利用（借阅与审批）")

add_h2("7.1 发起借阅（检索门户）")
add_num(1, "在检索结果或档案详情中把档案加入「借阅车」（电子到件、实体到卷）。")
add_num(2, "门户「我的借阅 → 借阅车」统一结算：逐件选择电子权限（在线浏览 / 下载 / 打印）或实体外借（原件 / 复印件），"
           "填写借阅事由与周期（最长 30 天）。")
add_num(3, "「审批链预览」按流程配置实时计算：基础链 →（含下载 / 打印 / 实体）升级 →（涉密）升级 → 终审。提交后进入审批。")

add_h2("7.2 审批（菜单：档案利用 → 审批中心）")
add_bullet("待办按当前角色的审批节点过滤；通过 / 驳回须填意见；终审通过系统自动拆单履约。")
add_bullet("审批步骤角色在服务端按当前步骤严格校验（防越级审批，admin 不豁免）。")
add_bullet("审批链规则由「档案配置 → 流程配置 → 借阅利用」的组链规则驱动，管理员修改后对今后发起的申请生效（在途单按原链）。")

add_h2("7.3 履约与归还（菜单：档案利用 → 借阅管理 / 借阅台账）")
add_bullet("电子：授权单即时生效，门户「在线调阅」限时访问（到期自动收回）。")
add_bullet("实体：出库核销 → 借出 → 归还核销；逾期进入红黑榜并触发黑名单（未还前禁新借）。")
add_bullet("已借出的卷支持预约排队，归还后自动通知下一预约人。")
add_bullet("借阅台账：全生命周期记录（谁 / 何时 / 借什么 / 审批链 / 应还实还 / 状态）。")
add_para("借阅单状态机：审批中 →（终审通过）履约中 → 已归还 / 已收回；驳回 / 撤销（申请人本人）终止。"
         "履约子单十态流转（电子授权 granted / 实体出库 checkout / 借出 / 归还等），"
         "每日巡检 Job 自动检测逾期并触发黑名单。借阅单号由后端 ams_code_serial 原子取号。")

# ============ 第 8 章 档案处置 ============
add_h1("第 8 章 档案处置")

add_h2("8.1 档案打包（菜单：档案处置 → 档案打包）")
add_para("用途：案卷封装与移交前整理，生成标准封装包（含元数据与版式文件清单）。"
         "数据源为全量件（收集池 ∪ 案卷库 ∪ 盒库），数据到达自动重建封装单元。")

add_h2("8.2 档案移交（菜单：档案处置 → 档案移交 / 档案利用 → 案卷移交管理）")
add_para("用途：会计部 → 档案部移交（临时保管期满）与对外移交的批次管理。后端 /transfers 提供批次 CRUD 与状态机"
         "（pending → prepared → received）+ 退回 / 删除（仅 pending）+ 卷明细实时解析 + 清册打印 + 每步写操作日志。", space_after=3)
add_bullet("案卷移交管理（TransferManagePage）：已入库卷勾选 → 发起移交弹窗 → 批次记录（生成清册 / 签收 / 退回 / 删除 / 清册打印）。")
add_bullet("档案移交（ArchiveTransferPage）：三栏批次执行看板（待准备 / 待签收 / 已移交），与移交管理页同一份台账。")
add_para("语义边界：移交归盒（卷 → 盒库，所内归档）≠ 对外移交（批次台账，卷节点状态不动）。"
         "创建校验：案卷须为已入库（transferred）才可对外移交。")

add_h2("8.3 鉴定销毁（菜单：档案处置 → 鉴定销毁）")
add_para("用途：档案到期鉴定与销毁执行，全环节真实状态机流转并写入不可篡改操作日志（全程留痕可审计）。", space_after=3)
add_num(1, "期满测算：按保管期限（10 年 / 30 年 / 永久）实时扫描盒库已入库卷，按「（年度 + 期限 + 1）-01-01」算期满日（永久不期满）。")
add_num(2, "鉴定评审：一键登记扫描到期卷为待鉴定任务（幂等），逐卷评审（销毁 / 留存，意见留痕）。")
add_num(3, "销毁执行：评审通过后执行销毁，删除 Alfresco 卷节点（级联卷内件）+ 盒计数回退 + 记录销毁时间 + 操作日志。")
add_para("页面组织：到期测算表 / 待鉴定评审 / 待销毁执行 / 已办结四区块。")

# ============ 第 9 章 档案统计 ============
add_h1("第 9 章 档案统计")
add_table(
    ["页面", "内容"],
    [
        ["统计驾驶舱", "库藏 / 流程 / 利用 / 合规一屏总览（12 个可配置模块，默认首页）"],
        ["库藏统计", "按类型 / 年度 / 期限 / 全宗 / 部门 / 载体家底盘点"],
        ["流程统计", "归档 / 组卷 / 四性检测 / 移交 / 鉴定处置全生命周期监控"],
        ["借阅统计", "借阅热力 / 逾期红黑榜 / 全链路操作日志"],
        ["合规统计", "期限 / 数据质量 / 安全 / 审计支撑（79 号令 + DA/T 94）"],
    ],
    col_widths=[3.0, 10.5]
)
add_para("说明：统计数据基于真实后端数据实时计算（非 mock），件级统计口径为全量件（收集池 ∪ 案卷库 ∪ 盒库），"
         "流程域移交批次 / 件数由已移交卷实时推导。驾驶舱模块的开关 / 排序 / 布局在「档案配置 → 驾驶舱配置」维护。")

# ============ 第 10 章 档案配置 ============
add_h1("第 10 章 档案配置")

add_h2("10.1 全宗管理 / 目录配置 / 元数据配置")
add_bullet("全宗管理：全宗（核算主体）档案存储总览与维护，真实 CRUD 于 Alfresco finance:fonds 节点。")
add_bullet("目录配置：多维业务科目档案目录分类体系，落 ams_config(directory)，驱动侧边栏 / 财务 / 项目视图。")
add_bullet("元数据配置：各档案门类元数据方案（字段 / 显示 / 上下文，26 列注册表一处写多处读），驱动全系统动态列与详情展示。")

add_h2("10.2 档号规则配置")
add_para("档号结构：全宗号 - KU·类别号·年度 - 保管期限 - 案卷号 - 件号（DA/T 13-2022）。"
         "服务端赋号引擎（VolumeService.assignOnConfirm / buildVolumeCode）实时读取 ams_config(archive-code-config)"
         "（赋号时机 / 流水位数 / 分隔符 / 类别前缀），写路径已贯通；流水号由 ams_code_serial 原子取号，彻底解决重号。"
         "页内分「档号规则定义 / 赋号时机 / 刚性规则 / 标准结构 / 规范溯源 / 电子专项 / 合规红线」七个分区。")

add_h2("10.3 档案三合一表配置")
add_para("分类体系 · 归档范围 · 保管期限三合一维护（79 号令），为法定口径只读标准展示页（非编辑配置）；"
         "是推送缺省期限推断与鉴定销毁测算的依据。页面标注「法定口径 · 只读标准 · 已生效」。")

add_h2("10.4 组卷盒号配置")
add_para("组卷规则（按类别独立规则：凭证 / 账簿 / 报告 / 其他）与盒号定义，落 ams_config(volume-grouping-config)；"
         "直接驱动组卷工作台「智能组卷」推荐（按类别 + 期间 + 连续号段分组）。")

add_h2("10.5 四性检测配置")
add_bullet("检测项库：环节（归档 / 移交 / 长期保存）× 四性 × 检测项的标准库，勾选集合即本单位检测方案；每项标注 DA/T · GB/T 依据。")
add_bullet("方案模板与四维配置：必填元数据字段、格式白名单、敏感词表、维度开关；保存即落 ams_config(inspection.plan)。")
add_bullet("服务端检测引擎实时消费方案：必填字段中文标签映射节点属性逐项核查、格式白名单 token→mime 映射、"
           "敏感关键词扫描；支持 POST /inspection/run-batch 对收集池批量检测，页面「立即执行检测」显示真实结果。")
add_bullet("检测报告逐条落库（四性状态 + 问题明细），支持人工复检留痕（复检人 / 原因 / 时间）。")

add_h2("10.6 报告配置 / 水印配置")
add_bullet("报告配置：报表类档案的报告模板与输出规则，整页重写为真实 CRUD 登记册，落 ams_config(report.config)（新增 / 启停 / 删除 / 检索）。")
add_bullet("水印配置：安全溯源水印策略，预览 / 下载 / 打印三场景动态水印，落 ams_config(watermark-config-v1) 并由服务端 PDFBox 真实烧录。")

add_h2("10.7 流程配置")
add_para("可视化流程设计器（画布编排）+ 业务元数据。「借阅利用」流程的「审批组链规则」为运行时真消费："
         "基础链（可排序增删角色）→ 升级规则（含下载 / 打印 / 实体、涉密两条）→ 终审角色；修改即时生效于今后发起的借阅申请，"
         "停用回退系统默认链。当前借阅审批路由由服务端内置执行，设计器为规则登记册 + 未来引擎底座（页面如实标注执行口径）。")

add_h2("10.8 库房配置")
add_para("库房布局全配置化：库房（增 / 删 / 改名）→ 密集架（增 / 删 / 改名 / 改维度：列 × 层 × 每层盒位）。", space_after=3)
add_bullet("库房号 / 架号创建后不可改（被架位引用），名称随时可改。")
add_bullet("删库房须无架；删架须无在架盒；缩小架维度时新边界外不得有在架盒（服务端强校验）；扩容自由。")
add_bullet("实体库房页按布局动态渲染，「自动上架」按配置顺序分配格位。")

add_h2("10.9 驾驶舱配置")
add_para("统计驾驶舱 12 个模块的开关、排序与布局管理，落 ams_config(cockpit) 由统计驾驶舱真实消费。")

# ============ 第 11 章 系统管理 ============
add_h1("第 11 章 系统管理")

add_h2("11.1 组织与权限")
add_bullet("单位管理 / 组织管理 / 人员管理：统一组织层级、部门与系统用户备案，真实 CRUD 于 Alfresco Groups / People + ams_user_ext。")
add_bullet("角色管理：业务角色划分与菜单权限矩阵（决定各角色可见菜单与后台入口），落 Alfresco ROLE_* 组 + ams_config(role-menus-v3)。")

add_h2("11.2 连接配置（菜单：系统管理 → 连接配置）")
add_para("收敛全部抓取 / 推送相关配置，三 Tab 统一管理（仅档案管理员 / 主管 / admin 可读写）：", space_after=3)
add_bullet("数据源连接：多数据源（用友 BIP / 金蝶 / 发票 / 银行 / 报销等）连接配置，每源含抓取计划（启用 / cron）"
           "与默认去向，secret 脱敏回显；用友支持「测试连接」实调网关验证。")
add_bullet("推送接入应用：签发 AppKey / AppSecret（仅展示一次）、接入应用列表、默认去向（签发时选择 + 列表行内切换）、推送批次历史 + 明细。")
add_bullet("接口字段映射：低代码字段映射（来源系统列表、规则编辑：类别 / 标准字段注册表 / 来源路径 / 7 种转换规则 / 默认值、样例 JSON 试映射）。")

add_h2("11.3 安全审计日志")
add_para("安全通道全链路行为追溯（ams_operation_log 哈希链防篡改）。操作日志 append-only + SHA-256 链 + "
         "数据库触发器禁止改删；后端提供 GET /audit/logs（读取）与 GET /audit/verify（真实验链，逐环重算哈希，"
         "如实区分通过 / 历史不可验 / 断链）。系统管理页的审计日志面板为真实数据。")

# ============ 第 12 章 附录 ============
add_h1("第 12 章 附录")

add_h2("12.1 档号结构速查")
add_para("Z001-KU·01·2026-D30-B001-0003-0012 = 全宗 Z001 · 会计门类 KU · 凭证类 01 · 2026 年 · 30 年期限 · "
         "盒流水 B001 · 卷流水 0003 · 件号 0012。")

add_h2("12.2 核心状态机速查")
add_table(
    ["域", "状态流"],
    [
        ["件", "仅件数据 → 待审核（入草稿卷）→ 已组卷（确认）"],
        ["案卷", "草稿 → 已确认（赋号）→ 已移交（归盒）→（退回 / 撤销回草稿）"],
        ["盒", "装盒中 active → 已封盒 sealed → 在架 stored →（下架回 sealed）"],
        ["借阅单", "审批中 →（终审通过）履约中 → 已归还 / 已收回；驳回 / 撤销终止"],
        ["移交批次", "待准备 pending → 待签收 prepared → 已移交 received；退回 / 删除（仅 pending）"],
        ["鉴定销毁", "待鉴定 → 评审通过（销毁 / 留存）→ 执行销毁（已销毁留痕）"],
    ],
    col_widths=[2.5, 11.0]
)

add_h2("12.3 归档环节检测项速查（节选）")
add_table(
    ["编码", "检测项", "四性", "依据"],
    [
        ["GD-1-01", "电子文件存在性", "真实性", "DA/T 94-2022"],
        ["GD-1-03", "档号规范性", "真实性", "DA/T 13-2022"],
        ["GD-2-01", "必填元数据齐全", "完整性", "DA/T 94-2022 / 79 号令"],
        ["GD-2-02", "凭证号连续性（断号）", "完整性", "DA/T 42-2022"],
        ["GD-2-03", "凭证号重复性", "完整性", "79 号令"],
        ["GD-3-01", "格式合规（白名单）", "可用性", "GB/T 33190-2016"],
        ["GD-4-01", "敏感信息模式扫描", "安全性", "—"],
    ],
    col_widths=[2.0, 5.5, 2.5, 3.5]
)
add_para("完整清单见系统内「档案配置 → 四性检测配置 → 检测项库」（含移交 / 长期保存环节）。")

add_h2("12.4 常见问题")
add_bullet("页面提示「服务无响应 / 请求超时」：后端 ams-server 未启动或版本过旧，请按 12.5 重启。")
add_bullet("看不到某菜单：角色无权限，联系系统管理员在「角色管理」调整菜单矩阵（新菜单如「鉴定销毁」需在角色矩阵手动开启）。")
add_bullet("确认组卷不可点：四性检测未通过或未运行，按案卷卡片问题明细处理后重新检测。")
add_bullet("架位选不中：该格已被占用或该架维度已在库房配置中调整，刷新后重试。")
add_bullet("新推送 / 抓取的数据在核对台看不到：核对 / 组卷池需刷新，系统会在采集成功后自动刷新件域镜像。")

add_h2("12.5 服务运维速查")
add_para("后端（ams-server）启动：")
add_para("cd ams-server && AMS_DB_URL=\"jdbc:postgresql://localhost:5432/alfresco?currentSchema=ams\" "
         "java -jar target/ams-server-0.0.1-SNAPSHOT.jar", indent=True, font_cn="Consolas")
add_para("前端：financeWeb 目录 npm run dev（:5000）；数据库：PostgreSQL 16.5（docker，:5432 / :15432，schema=ams）；"
         "内容库：Alfresco Community（:8080，finance-model 自定义模型需按规程热部署）。"
         "前端经 /api/ams 代理至 ams-server（:8081）、/api/proxy/alfresco 至 Alfresco（:8080）。")

# ============ 保存 ============
OUT = "/tmp/manual_gen/会计档案管理系统操作手册.docx"
DOC.save(OUT)
print("已生成:", OUT)
